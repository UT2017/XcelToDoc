#!/usr/bin/env python
# coding: utf-8

# In[ ]:


import platform
import tkinter as tk
import openpyxl as xl
import sys
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl.utils import get_column_letter, column_index_from_string
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE

from docx.shared import Pt
from docx.shared import RGBColor

from docx.enum.style import WD_STYLE_TYPE


# In[ ]:


label_count = 11


def update_doc_helper():
    global label_count

    # keeps track of success or failure
    var1 = 0

    try:
        update_doc()

    except Exception as ex:

        # debug info
        exc_type, exc_obj, exc_tb = sys.exc_info()
        fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        print(exc_type, fname, exc_tb.tb_lineno)

        message1 = (
            "\nException type "
            + str(exc_type)
            + "\nFile name "
            + str(fname)
            + "\nLine no "
            + str(exc_tb.tb_lineno)
        )

        template = "An exception of type {0} occurred. Arguments:\n{1!r}"
        message = template.format(type(ex).__name__, ex.args)

        tk.Label(
            frame,
            anchor=tk.W,
            justify=tk.LEFT,
            font="Helvetica 20 bold",
            foreground="red",
            bg="#ffffff",
            text=message + "\n",
        ).grid(row=label_count, sticky=tk.W, padx=(4, 4))
        label_count = label_count + 1

        tk.Label(
            frame,
            anchor=tk.W,
            justify=tk.LEFT,
            font="Helvetica 20 bold",
            foreground="red",
            bg="#ffffff",
            text=message1 + "\n",
        ).grid(row=label_count, sticky=tk.W, padx=(4, 4))
        label_count = label_count + 1

        var1 = 1

    finally:
        if var1 == 0:
            tk.Label(
                frame,
                anchor=tk.W,
                justify=tk.LEFT,
                font="Helvetica 20 bold",
                foreground="green",
                bg="#ffffff",
                text="Success!\nPlease find updated doc at " + os.getcwd(),
            ).grid(row=label_count, sticky=tk.W, padx=(4, 4))
            label_count = label_count + 1


def update_doc():

    doc = Document()

    input_master_excel_name_xlsx = input_master_excel_name.get() + ".xlsx"
    output_doc_name_docx = output_doc_name.get() + ".docx"

    wb = xl.load_workbook(input_master_excel_name_xlsx)
    res = len(wb.sheetnames)

    def add_para(content):
        font_name = "Arial"
        font_size = 11
        paragraph = doc.add_paragraph(content)
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)

    def add_heading(content1):
        font_name = "Arial"
        font_size = 14
        para = doc.add_paragraph()
        runner = para.add_run(content1)

        font1 = runner.font
        font1.name = font_name
        font1.size = Pt(font_size)

    # keeps track of number of sheets
    varlen = 1

    # iterate over sheet by sheet
    for sheet_iterator in wb:

        # get a list of all rows
        sheet_rows_list = list(sheet_iterator.rows)

        print_table = 0

        row_iterator = 1
        col_iterator = 1

        # iterate over array of rows
        for sheet_rows_iterator in sheet_rows_list:

            print_para = 0
            print_heading = 0

            col_iterator = 1

            if print_table == 0:

                printwhichrow = 0
                printwhichcol = 0

                max_col = sys.maxsize
                max_row = sys.maxsize

                start_row = 0
                start_col = 0

                current_row = 0
                current_col = 0

                firstrowdone = 0

                keepaddingrows = 0

                lastrow = 0

            if print_table == 1 and keepaddingrows == 1:
                table.add_row()
                current_row = current_row + 1

            # iterate over cells in a row
            for sheet_row_iterator in sheet_rows_iterator:
                cell_value = sheet_row_iterator.value

                if cell_value is None:
                    cell_value = ""

                # handling tag "TRE"
                if keepaddingrows == 1 and cell_value == "TRE":
                    keepaddingrows = 0
                    firstrowdone = 0
                    lastrow = 1

                elif (
                    print_table == 1
                    and lastrow == 1
                    and current_row < max_row
                    and current_col < max_col
                    and row_iterator >= start_row
                    and row_iterator < (start_row + max_row)
                    and col_iterator >= start_col
                    and col_iterator < (start_col + max_col)
                ):

                    table.cell(current_row, current_col).text = str(cell_value)
                    current_col = current_col + 1

                # table ends
                elif (
                    print_table == 1
                    and lastrow == 1
                    and current_row < max_row
                    and current_col == max_col
                    and col_iterator == (start_col + max_col)
                ):

                    print_table = 0
                    lastrow = 0
                    start_row = 0
                    start_col = 0
                    current_row = 0
                    current_col = 0

                    # keeps track if first row
                    var11 = 0

                    # formatting the table printed

                    # for each table cell, add font
                    for row in table.rows:
                        for cell in row.cells:

                            # set each cell width as auto
                            # still doesn't span to page width; need to search a method
                            cell._tc.tcPr.tcW.type = "auto"

                            paragraphs = cell.paragraphs

                            # if first line of table, align to center
                            # center alignment done to each paragraph present in a cell
                            if var11 == 0:
                                for paragraph in paragraphs:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                            # changing font for each paragraph present in a cell through their runner
                            for run in paragraph.runs:
                                font = run.font
                                font.name = "Arial"
                                font.size = Pt(10)
                                # if first line of table, make text to bold
                                if var11 == 0:
                                    font.bold = True

                        # after first row, var11 becomes 1
                        var11 = 1

                    # add an empty line after each table
                    add_para("\n")

                # handling tag "TCE"
                elif print_table == 1 and firstrowdone == 0 and cell_value == "TCE":
                    firstrowdone = 1
                    max_col = current_col
                    current_col = 0

                elif (
                    print_table == 1
                    and firstrowdone == 0
                    and current_row < max_row
                    and current_col < max_col
                    and row_iterator >= start_row
                    and row_iterator < (start_row + max_row)
                    and col_iterator >= start_col
                    and col_iterator < (start_col + max_col)
                ):
                    a = 367284
                    table.add_column(a)
                    table.cell(current_row, current_col).text = str(cell_value)

                    current_col = current_col + 1

                elif (
                    print_table == 1
                    and firstrowdone == 1
                    and current_row < max_row
                    and current_col < max_col
                    and row_iterator >= start_row
                    and row_iterator < (start_row + max_row)
                    and col_iterator >= start_col
                    and col_iterator < (start_col + max_col)
                ):

                    table.cell(current_row, current_col).text = str(cell_value)
                    current_col = current_col + 1

                elif (
                    print_table == 1
                    and firstrowdone == 1
                    and current_row < max_row
                    and current_col == max_col
                    and col_iterator == (start_col + max_col)
                ):
                    current_col = 0

                elif print_para == 1:
                    add_para(cell_value)
                    print_para = 0

                elif print_heading == 1:
                    add_heading(cell_value)
                    print_heading = 0

                # handling tag "P"
                elif cell_value == "P":
                    print_para = 1

                # handling tag "H"
                elif cell_value == "H":
                    print_heading = 1

                # handling tag "TS"
                elif cell_value == "TS":
                    print_table = 1
                    start_row = int(sheet_row_iterator.row)
                    start_col = (
                        int(column_index_from_string(sheet_row_iterator.column)) + 1
                    )

                    table = doc.add_table(1, 0)
                    table.style = "TableGrid"
                    table.allow_autofit = True

                    firstrowdone = 0
                    keepaddingrows = 1

                    current_row = 0
                    current_col = 0

                    printwhichrow = start_row
                    printwhichcol = start_col

                col_iterator = col_iterator + 1

            row_iterator = row_iterator + 1

        # add a page break after each sheet ends
        # if sheet is not last sheet then add page break
        if varlen != res:
            para = doc.add_paragraph("")
            run = para.add_run()
            run.add_break(WD_BREAK.PAGE)

        varlen = varlen + 1

    doc.save(output_doc_name_docx)


# In[ ]:


"""GUI"""


# In[ ]:


main_window = tk.Tk()
main_window.configure(background="#ffffff")
main_window.title("XcelToDoc")


# In[ ]:


def onFrameConfigure(canvas):
    """Reset the scroll region to encompass the inner frame"""
    canvas.configure(scrollregion=canvas.bbox("all"))


# In[ ]:


def _on_mousewheel(event):
    canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")


# In[ ]:


canvas = tk.Canvas(main_window, borderwidth=0, bg="#ffffff")
frame = tk.Frame(canvas, bg="#ffffff")

# scrollbar
vsb = tk.Scrollbar(main_window, orient="vertical", command=canvas.yview)
canvas.configure(yscrollcommand=vsb.set)
hsb = tk.Scrollbar(main_window, orient="horizontal", command=canvas.xview)
canvas.configure(xscrollcommand=hsb.set)
canvas.bind_all("<MouseWheel>", _on_mousewheel)

hsb.pack(side="bottom", fill="x")
vsb.pack(side="right", fill="y")
canvas.pack(side="left", fill="both", expand=True)
canvas.create_window((0, 0), window=frame, anchor="nw")

frame.bind("<Configure>", lambda event, canvas=canvas: onFrameConfigure(canvas))


# In[ ]:


tk.Label(
    frame,
    anchor=tk.W,
    justify=tk.LEFT,
    font="Helvetica 20 bold",
    foreground="#ffffff",
    bg="#ff7919",
    text="Excel to Doc made easy",
).grid(sticky=(tk.N, tk.S, tk.E, tk.W), rowspan=2, columnspan=3)

# Printing 1 empty row
tk.Label(frame, anchor=tk.W, justify=tk.LEFT, bg="#ffffff").grid(
    row=2, sticky=(tk.N, tk.S, tk.E, tk.W)
)

tk.Label(
    frame,
    anchor=tk.W,
    justify=tk.LEFT,
    font="Helvetica 12 ",
    bg="#ffffff",
    text="Input Master excel file name without the extension(.xlsx)",
).grid(row=3, sticky=(tk.N, tk.S, tk.E, tk.W), padx=(4, 4))
tk.Label(
    frame,
    anchor=tk.W,
    justify=tk.LEFT,
    font="Helvetica 12 ",
    bg="#ffffff",
    text="Input updated (output) Doc file name without the extension(.docx)",
).grid(row=4, sticky=(tk.N, tk.S, tk.E, tk.W), padx=(4, 4))

# Printing 3 empty rows
tk.Label(frame, anchor=tk.W, justify=tk.LEFT, bg="#ffffff").grid(
    row=6, sticky=(tk.N, tk.S, tk.E, tk.W)
)
tk.Label(frame, anchor=tk.W, justify=tk.LEFT, bg="#ffffff").grid(
    row=8, sticky=(tk.N, tk.S, tk.E, tk.W)
)
tk.Label(frame, anchor=tk.W, justify=tk.LEFT, bg="#ffffff").grid(
    row=9, sticky=(tk.N, tk.S, tk.E, tk.W)
)

frame.columnconfigure(0, weight=100, minsize=700)
frame.columnconfigure(1, weight=50, minsize=600)
frame.columnconfigure(2, weight=3000, minsize=600)

input_master_excel_name = tk.Entry(
    frame, selectborderwidth=100, relief="sunken", width=60
)
output_doc_name = tk.Entry(frame, selectborderwidth=100, relief="sunken", width=60)

input_master_excel_name.grid(row=3, column=1, sticky=tk.W)
output_doc_name.grid(row=4, column=1, sticky=tk.W)


if platform.system() == "Darwin":  # if its a Mac
    tk.Button(
        frame,
        font="Helvetica 12",
        width=30,
        highlightbackground="#3E4149",
        bg="#3B5998",
        fg="#ffffff",
        text="Get updated doc",
        command=update_doc_helper,
        justify=tk.LEFT,
    ).grid(row=7, column=0, sticky=tk.W, padx=(4, 4))
else:  # if its Windows or Linux
    tk.Button(
        frame,
        font="Helvetica 12",
        width=30,
        bg="#3B5998",
        fg="#ffffff",
        text="Get updated doc",
        command=update_doc_helper,
        justify=tk.LEFT,
    ).grid(row=7, column=0, sticky=tk.W, padx=(4, 4))
    tk.Button(
        frame,
        font="Helvetica 12",
        width=30,
        bg="#3B5998",
        fg="#ffffff",
        text="Quit",
        command=main_window.destroy,
        justify=tk.LEFT,
    ).grid(row=7, column=1, sticky=tk.W, padx=(4, 4))


main_window.mainloop()


# In[ ]:
